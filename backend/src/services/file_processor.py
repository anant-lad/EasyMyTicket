"""
File Processor Service
Extracts text and data from various file formats for ticket context generation
"""
import os
import json
from typing import Dict, Any, Optional, List
from pathlib import Path
import traceback


class FileProcessor:
    """
    Processes uploaded files to extract text, tables, and other structured data
    Supports: PDF, DOCX, TXT, XML, HTML, CSV, XLSX, and images (with OCR)
    """
    
    def __init__(self, enable_ocr: bool = True, ocr_engine: str = 'pytesseract'):
        """
        Initialize file processor
        
        Args:
            enable_ocr: Whether to enable OCR for images
            ocr_engine: OCR engine to use ('pytesseract' or 'easyocr')
        """
        self.enable_ocr = enable_ocr
        self.ocr_engine = ocr_engine
        self._lazy_imports = {}
    
    def _import_pdf_libraries(self):
        """Lazy import PDF processing libraries"""
        if 'pdf' not in self._lazy_imports:
            try:
                import pdfplumber
                self._lazy_imports['pdfplumber'] = pdfplumber
            except ImportError:
                print("Warning: pdfplumber not installed, falling back to PyPDF2")
                try:
                    import PyPDF2
                    self._lazy_imports['PyPDF2'] = PyPDF2
                except ImportError:
                    print("Error: No PDF library available")
                    self._lazy_imports['pdf'] = None
    
    def _import_docx_libraries(self):
        """Lazy import DOCX processing libraries"""
        if 'docx' not in self._lazy_imports:
            try:
                import docx
                self._lazy_imports['docx'] = docx
            except ImportError:
                print("Error: python-docx not installed")
                self._lazy_imports['docx'] = None
    
    def _import_excel_libraries(self):
        """Lazy import Excel/CSV processing libraries"""
        if 'pandas' not in self._lazy_imports:
            try:
                import pandas as pd
                self._lazy_imports['pandas'] = pd
            except ImportError:
                print("Error: pandas not installed")
                self._lazy_imports['pandas'] = None
    
    def _import_image_libraries(self):
        """Lazy import image processing and OCR libraries"""
        if 'image' not in self._lazy_imports:
            try:
                from PIL import Image
                self._lazy_imports['PIL'] = Image
                
                if self.enable_ocr:
                    if self.ocr_engine == 'pytesseract':
                        import pytesseract
                        self._lazy_imports['pytesseract'] = pytesseract
                    elif self.ocr_engine == 'easyocr':
                        import easyocr
                        self._lazy_imports['easyocr'] = easyocr
            except ImportError as e:
                print(f"Error importing image libraries: {e}")
                self._lazy_imports['image'] = None
    
    def _import_html_libraries(self):
        """Lazy import HTML/XML processing libraries"""
        if 'bs4' not in self._lazy_imports:
            try:
                from bs4 import BeautifulSoup
                self._lazy_imports['bs4'] = BeautifulSoup
            except ImportError:
                print("Error: beautifulsoup4 not installed")
                self._lazy_imports['bs4'] = None
    
    def process_file(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """
        Process a file and extract its content
        
        Args:
            file_path: Path to the file
            file_type: File extension (pdf, docx, txt, etc.)
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        result = {
            'file_path': file_path,
            'file_type': file_type,
            'extracted_text': '',
            'tables': [],
            'images_analysis': [],
            'metadata': {},
            'processing_status': 'pending',
            'error': None
        }
        
        try:
            file_type = file_type.lower().strip('.')
            
            if file_type == 'pdf':
                result = self.extract_pdf(file_path)
            elif file_type in ['docx', 'doc']:
                result = self.extract_docx(file_path)
            elif file_type == 'txt':
                result = self.extract_txt(file_path)
            elif file_type in ['csv', 'xlsx', 'xls']:
                result = self.extract_csv_xlsx(file_path, file_type)
            elif file_type in ['xml', 'html']:
                result = self.extract_xml_html(file_path, file_type)
            elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
                result = self.extract_image(file_path)
            else:
                result['processing_status'] = 'unsupported'
                result['error'] = f'Unsupported file type: {file_type}'
            
            result['file_path'] = file_path
            result['file_type'] = file_type
            
        except Exception as e:
            result['processing_status'] = 'failed'
            result['error'] = str(e)
            print(f"Error processing file {file_path}: {e}")
            traceback.print_exc()
        
        return result
    
    def extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text and tables from PDF"""
        self._import_pdf_libraries()
        
        result = {
            'extracted_text': '',
            'tables': [],
            'metadata': {},
            'processing_status': 'completed'
        }
        
        try:
            # Try pdfplumber first (better table extraction)
            if 'pdfplumber' in self._lazy_imports and self._lazy_imports['pdfplumber']:
                pdfplumber = self._lazy_imports['pdfplumber']
                with pdfplumber.open(file_path) as pdf:
                    text_parts = []
                    for page_num, page in enumerate(pdf.pages, 1):
                        # Extract text
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                        
                        # Extract tables
                        tables = page.extract_tables()
                        if tables:
                            for table_idx, table in enumerate(tables):
                                result['tables'].append({
                                    'page': page_num,
                                    'table_index': table_idx,
                                    'data': table
                                })
                    
                    result['extracted_text'] = '\n\n'.join(text_parts)
                    result['metadata'] = {
                        'pages': len(pdf.pages),
                        'method': 'pdfplumber'
                    }
            
            # Fallback to PyPDF2
            elif 'PyPDF2' in self._lazy_imports and self._lazy_imports['PyPDF2']:
                PyPDF2 = self._lazy_imports['PyPDF2']
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text_parts = []
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        text = page.extract_text()
                        if text:
                            text_parts.append(f"--- Page {page_num} ---\n{text}")
                    
                    result['extracted_text'] = '\n\n'.join(text_parts)
                    result['metadata'] = {
                        'pages': len(pdf_reader.pages),
                        'method': 'PyPDF2'
                    }
            else:
                result['processing_status'] = 'failed'
                result['error'] = 'No PDF library available'
        
        except Exception as e:
            result['processing_status'] = 'failed'
            result['error'] = str(e)
            print(f"Error extracting PDF: {e}")
        
        return result
    
    def extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text and tables from DOCX"""
        self._import_docx_libraries()
        
        result = {
            'extracted_text': '',
            'tables': [],
            'metadata': {},
            'processing_status': 'completed'
        }
        
        try:
            if 'docx' in self._lazy_imports and self._lazy_imports['docx']:
                docx = self._lazy_imports['docx']
                doc = docx.Document(file_path)
                
                # Extract paragraphs
                text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
                result['extracted_text'] = '\n'.join(text_parts)
                
                # Extract tables
                for table_idx, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    
                    result['tables'].append({
                        'table_index': table_idx,
                        'data': table_data
                    })
                
                result['metadata'] = {
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'method': 'python-docx'
                }
            else:
                result['processing_status'] = 'failed'
                result['error'] = 'python-docx not available'
        
        except Exception as e:
            result['processing_status'] = 'failed'
            result['error'] = str(e)
            print(f"Error extracting DOCX: {e}")
        
        return result
    
    def extract_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file"""
        result = {
            'extracted_text': '',
            'tables': [],
            'metadata': {},
            'processing_status': 'completed'
        }
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                result['extracted_text'] = file.read()
                result['metadata'] = {
                    'encoding': 'utf-8',
                    'method': 'direct'
                }
        except Exception as e:
            result['processing_status'] = 'failed'
            result['error'] = str(e)
            print(f"Error extracting TXT: {e}")
        
        return result
    
    def extract_csv_xlsx(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Extract data from CSV or Excel files"""
        self._import_excel_libraries()
        
        result = {
            'extracted_text': '',
            'tables': [],
            'metadata': {},
            'processing_status': 'completed'
        }
        
        try:
            if 'pandas' in self._lazy_imports and self._lazy_imports['pandas']:
                pd = self._lazy_imports['pandas']
                
                # Read file based on type
                if file_type == 'csv':
                    df = pd.read_csv(file_path)
                else:  # xlsx or xls
                    df = pd.read_excel(file_path)
                
                # Convert to text representation
                result['extracted_text'] = df.to_string()
                
                # Store structured data
                result['tables'].append({
                    'table_index': 0,
                    'data': df.values.tolist(),
                    'columns': df.columns.tolist(),
                    'shape': df.shape
                })
                
                result['metadata'] = {
                    'rows': len(df),
                    'columns': len(df.columns),
                    'column_names': df.columns.tolist(),
                    'method': 'pandas'
                }
            else:
                result['processing_status'] = 'failed'
                result['error'] = 'pandas not available'
        
        except Exception as e:
            result['processing_status'] = 'failed'
            result['error'] = str(e)
            print(f"Error extracting CSV/XLSX: {e}")
        
        return result
    
    def extract_xml_html(self, file_path: str, file_type: str) -> Dict[str, Any]:
        """Extract text from XML or HTML files"""
        self._import_html_libraries()
        
        result = {
            'extracted_text': '',
            'tables': [],
            'metadata': {},
            'processing_status': 'completed'
        }
        
        try:
            if 'bs4' in self._lazy_imports and self._lazy_imports['bs4']:
                BeautifulSoup = self._lazy_imports['bs4']
                
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
                
                parser = 'html.parser' if file_type == 'html' else 'xml'
                soup = BeautifulSoup(content, parser)
                
                # Extract text
                result['extracted_text'] = soup.get_text(separator='\n', strip=True)
                
                result['metadata'] = {
                    'parser': parser,
                    'method': 'beautifulsoup4'
                }
            else:
                # Fallback: read as plain text
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    result['extracted_text'] = file.read()
                result['metadata'] = {'method': 'plain_text'}
        
        except Exception as e:
            result['processing_status'] = 'failed'
            result['error'] = str(e)
            print(f"Error extracting XML/HTML: {e}")
        
        return result
    
    def extract_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from images using OCR"""
        self._import_image_libraries()
        
        result = {
            'extracted_text': '',
            'tables': [],
            'images_analysis': [],
            'metadata': {},
            'processing_status': 'completed'
        }
        
        try:
            if not self.enable_ocr:
                result['processing_status'] = 'skipped'
                result['error'] = 'OCR disabled'
                return result
            
            if 'PIL' not in self._lazy_imports or not self._lazy_imports['PIL']:
                result['processing_status'] = 'failed'
                result['error'] = 'PIL not available'
                return result
            
            PIL_Image = self._lazy_imports['PIL']
            image = PIL_Image.open(file_path)
            
            # Get image metadata
            result['metadata'] = {
                'size': image.size,
                'mode': image.mode,
                'format': image.format
            }
            
            # Perform OCR
            if self.ocr_engine == 'pytesseract' and 'pytesseract' in self._lazy_imports:
                pytesseract = self._lazy_imports['pytesseract']
                text = pytesseract.image_to_string(image)
                result['extracted_text'] = text
                result['metadata']['ocr_engine'] = 'pytesseract'
            
            elif self.ocr_engine == 'easyocr' and 'easyocr' in self._lazy_imports:
                # EasyOCR requires initialization (expensive, so we'd cache it)
                # For now, fallback to pytesseract or skip
                result['processing_status'] = 'skipped'
                result['error'] = 'EasyOCR not implemented yet'
            
            else:
                result['processing_status'] = 'failed'
                result['error'] = f'OCR engine {self.ocr_engine} not available'
        
        except Exception as e:
            result['processing_status'] = 'failed'
            result['error'] = str(e)
            print(f"Error extracting image: {e}")
        
        return result


# Singleton instance
_file_processor = None

def get_file_processor(enable_ocr: bool = True, ocr_engine: str = 'pytesseract') -> FileProcessor:
    """Get or create file processor instance"""
    global _file_processor
    if _file_processor is None:
        _file_processor = FileProcessor(enable_ocr, ocr_engine)
    return _file_processor
